# This is a sample Python script.

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.
# from spire.doc import *

# Perform a Pubmed Search for huntington and date range i.e (huntington) AND (("2025/04/01"[Date - Publication] : "2025/06/30"[Date - Publication]))
# Save all results in PubMed format
# Save all results in CSV format
# have the Excel file of members handy




import pandas as pd
from docx import Document
from tkinter import filedialog as fd
from datetime import datetime
from tkinter import messagebox as mbox
from tkinter import simpledialog as sd
from unidecode import unidecode
import tkinter as tk

class PMatch:
    def __init__(self, name, pmid, authors, title, citation):
        self.name = [name]
        self.pmid = pmid
        self.authors = authors
        self.title = title
        self.citation = citation



def ProcessPubmed():
    hd_list = {}
    pubmedfile = fd.askopenfilename(title='Select PubMed File')

    with open(pubmedfile, 'r', encoding="utf8") as f:
        lines = f.readlines()
        currentpmid = ''
        for line in lines:
            if line.startswith('PMID'):
                currentpmid = line.split('- ')[1].rstrip()
                hd_list[currentpmid] = {}
                hd_list[currentpmid]['authors'] = []
                hd_list[currentpmid]['date'] = ''
            if line.startswith('FAU'):
                hd_list[currentpmid]['authors'].append(line.split('- ')[1].rstrip())
            if line.startswith('DP'):
                td = line.split('- ')[1].rstrip()
                hd_list[currentpmid]['date'] = td
    return hd_list

def ProcessMembers():
    memberlist = fd.askopenfilename(title='Select Member Excel File')
    data = pd.read_excel(memberlist)
    return data

def ProcessCSV():
    csvfile = fd.askopenfilename(title='Select PubMedCSV File')
    cites = pd.read_csv(csvfile, encoding='utf8')
    return cites

def normalize(s):
    return unidecode(s).strip()

def CompareNames(data, cites):
    rpublist = {}
    apublist = {}
    # counter = 0
    for cat, lname, fname in zip(data['Membership category'],data['Last name'], data['First name']):
        if isinstance(cat, str):
            if 'regular' not in cat and 'associated' not in cat:
                continue
        if isinstance(lname, str):
            if 'http' in lname or 'http' in fname:
                continue
            for paper in hd_list.keys():
                if len(paper) > 0:
                    for author in hd_list[paper]['authors']:
                        if author.split(', ')[0].rstrip().lower() == lname.rstrip().lower():
                            if author.split(', ')[1].rstrip().split(' ')[0].lower() == fname.rstrip().lower():
                                temp = cites[cites.PMID == int(paper)][["Authors", "Title", "Citation"]].values
                                if cat == 'regular':
                                    if paper in rpublist.keys():
                                        rpublist[paper].name.append(lname.rstrip().lower() + ' ' + fname.rstrip().lower())
                                    else:
                                        rpublist[paper] = PMatch(lname.rstrip().lower() + ' ' + fname.rstrip().lower(),paper, temp[0][0], temp[0][1], temp[0][2] )
                                    # rpublist.add(lname.rstrip().lower() + '\t' + fname.rstrip().lower() + '\t' + temp[0][0] + '\t' + temp[0][1] + ' ' + temp[0][2] + ' PMID:' + paper)
                                if cat == 'associated':
                                    if paper in apublist.keys():
                                        apublist[paper].name.append(
                                            lname.rstrip().lower() + ' ' + fname.rstrip().lower())
                                    else:
                                        apublist[paper] = PMatch(lname.rstrip().lower() + ' ' + fname.rstrip().lower(),
                                                                 paper, temp[0][0], temp[0][1], temp[0][2])
                                    # apublist.add(lname.rstrip().lower() + '\t' + fname.rstrip().lower() + '\t' + temp[0][0] + '\t' + temp[0][1] + ' ' + temp[0][2] + ' PMID:' + paper)
                            elif len(author.split(', ')[1].rstrip().split(' ')[0].lower()) == 1 and author.split(', ')[1].rstrip().split(' ')[0].lower() == fname.rstrip()[0].lower():
                                temp = cites[cites.PMID == int(paper)][["Authors", "Title", "Citation"]].values
                                if cat == 'regular':
                                    if paper in rpublist.keys():
                                        rpublist[paper].name.append(lname.rstrip().lower() + ' ' + fname.rstrip().lower())
                                    else:
                                        rpublist[paper] = PMatch(lname.rstrip().lower() + ' ' + fname.rstrip().lower(),paper, temp[0][0], temp[0][1], temp[0][2] )                                    # rpublist.add(lname.rstrip().lower() + '\t' + fname.rstrip().lower() + '\t' + temp[0][0] + '\t' + temp[0][1] + ' ' + temp[0][2] + ' PMID:' + paper)
                                if cat == 'associated':
                                    if paper in apublist.keys():
                                        apublist[paper].name.append(lname.rstrip().lower() + ' ' + fname.rstrip().lower())
                                    else:
                                        apublist[paper] = PMatch(lname.rstrip().lower() + ' ' + fname.rstrip().lower(),paper, temp[0][0], temp[0][1], temp[0][2] )
    return rpublist, apublist

def CompareNamesN(data, cites):
    rpublist = {}
    apublist = {}
    # counter = 0
    for cat, lname, fname in zip(data['Membership category'],data['Last name'], data['First name']):
        if isinstance(cat, str):
            if 'regular' not in cat and 'associated' not in cat:
                continue
        if isinstance(lname, str):
            if 'http' in lname or 'http' in fname:
                continue
            for paper in hd_list.keys():
                if len(paper) > 0:
                    match = False
                    for author in hd_list[paper]['authors']:
                        if author.split(', ')[0].rstrip().lower() == lname.rstrip().lower():
                            if author.split(', ')[1].rstrip().split(' ')[0].lower() == fname.rstrip().lower():
                                match = True
                                temp = cites[cites.PMID == int(paper)][["Authors", "Title", "Citation"]].values
                                if cat == 'regular':
                                    if paper in rpublist.keys():
                                        rpublist[paper].name.append(lname.rstrip().lower() + ' ' + fname.rstrip().lower())
                                    else:
                                        rpublist[paper] = PMatch(lname.rstrip().lower() + ' ' + fname.rstrip().lower(),paper, temp[0][0], temp[0][1], temp[0][2] )
                                    # rpublist.add(lname.rstrip().lower() + '\t' + fname.rstrip().lower() + '\t' + temp[0][0] + '\t' + temp[0][1] + ' ' + temp[0][2] + ' PMID:' + paper)
                                if cat == 'associated':
                                    if paper in apublist.keys():
                                        apublist[paper].name.append(
                                            lname.rstrip().lower() + ' ' + fname.rstrip().lower())
                                    else:
                                        apublist[paper] = PMatch(lname.rstrip().lower() + ' ' + fname.rstrip().lower(),
                                                                 paper, temp[0][0], temp[0][1], temp[0][2])
                                    # apublist.add(lname.rstrip().lower() + '\t' + fname.rstrip().lower() + '\t' + temp[0][0] + '\t' + temp[0][1] + ' ' + temp[0][2] + ' PMID:' + paper)
                            elif len(author.split(', ')[1].rstrip().split(' ')[0].lower()) == 1 and author.split(', ')[1].rstrip().split(' ')[0].lower() == fname.rstrip()[0].lower():
                                match = True
                                temp = cites[cites.PMID == int(paper)][["Authors", "Title", "Citation"]].values
                                if cat == 'regular':
                                    if paper in rpublist.keys():
                                        rpublist[paper].name.append(lname.rstrip().lower() + ' ' + fname.rstrip().lower())
                                    else:
                                        rpublist[paper] = PMatch(lname.rstrip().lower() + ' ' + fname.rstrip().lower(),paper, temp[0][0], temp[0][1], temp[0][2] )                                    # rpublist.add(lname.rstrip().lower() + '\t' + fname.rstrip().lower() + '\t' + temp[0][0] + '\t' + temp[0][1] + ' ' + temp[0][2] + ' PMID:' + paper)
                                if cat == 'associated':
                                    if paper in apublist.keys():
                                        apublist[paper].name.append(lname.rstrip().lower() + ' ' + fname.rstrip().lower())
                                    else:
                                        apublist[paper] = PMatch(lname.rstrip().lower() + ' ' + fname.rstrip().lower(),paper, temp[0][0], temp[0][1], temp[0][2] )
                    if match == False:
                        for author in hd_list[paper]['authors']:
                            if normalize(author.split(', ')[0]).rstrip().lower() == lname.rstrip().lower():
                                if normalize(author.split(', ')[1]).rstrip().split(' ')[0].lower() == fname.rstrip().lower():
                                    match = True
                                    temp = cites[cites.PMID == int(paper)][["Authors", "Title", "Citation"]].values
                                    if cat == 'regular':
                                        if paper in rpublist.keys():
                                            rpublist[paper].name.append(lname.rstrip().lower() + ' ' + fname.rstrip().lower())
                                        else:
                                            rpublist[paper] = PMatch(lname.rstrip().lower() + ' ' + fname.rstrip().lower(),paper, temp[0][0], temp[0][1], temp[0][2] )
                                        # rpublist.add(lname.rstrip().lower() + '\t' + fname.rstrip().lower() + '\t' + temp[0][0] + '\t' + temp[0][1] + ' ' + temp[0][2] + ' PMID:' + paper)
                                    if cat == 'associated':
                                        if paper in apublist.keys():
                                            apublist[paper].name.append(
                                                lname.rstrip().lower() + ' ' + fname.rstrip().lower())
                                        else:
                                            apublist[paper] = PMatch(lname.rstrip().lower() + ' ' + fname.rstrip().lower(),
                                                                     paper, temp[0][0], temp[0][1], temp[0][2])
                                        # apublist.add(lname.rstrip().lower() + '\t' + fname.rstrip().lower() + '\t' + temp[0][0] + '\t' + temp[0][1] + ' ' + temp[0][2] + ' PMID:' + paper)
                                elif len(author.split(', ')[1].rstrip().split(' ')[0].lower()) == 1 and normalize(author.split(', ')[1]).rstrip().split(' ')[0].lower() == fname.rstrip()[0].lower():
                                    match = True
                                    temp = cites[cites.PMID == int(paper)][["Authors", "Title", "Citation"]].values
                                    if cat == 'regular':
                                        if paper in rpublist.keys():
                                            rpublist[paper].name.append(lname.rstrip().lower() + ' ' + fname.rstrip().lower())
                                        else:
                                            rpublist[paper] = PMatch(lname.rstrip().lower() + ' ' + fname.rstrip().lower(),paper, temp[0][0], temp[0][1], temp[0][2] )                                    # rpublist.add(lname.rstrip().lower() + '\t' + fname.rstrip().lower() + '\t' + temp[0][0] + '\t' + temp[0][1] + ' ' + temp[0][2] + ' PMID:' + paper)
                                    if cat == 'associated':
                                        if paper in apublist.keys():
                                            apublist[paper].name.append(lname.rstrip().lower() + ' ' + fname.rstrip().lower())
                                        else:
                                            apublist[paper] = PMatch(lname.rstrip().lower() + ' ' + fname.rstrip().lower(),paper, temp[0][0], temp[0][1], temp[0][2] )
    return rpublist, apublist

# Press the green button in the gutter to run the script.
if __name__ == '__main__':

    hd_list = None
    data = None
    cites = None

    welcometext = ['To use the EHDN PubMed Parser you will need three files: a PubMed Text File, a PubMed CSV File, an Excel file of the membership \n',
                   'Creating the PubMed Files:\n ',
                   'Perform a Pubmed Search for huntington and a date range i.e (huntington) AND (("2025/04/01"[Date - Publication] : "2025/06/30"[Date - Publication]))\n',
                   'Save all results in PubMed format',
                   'Save all results in CSV format\n',
                   'Please select the files in the following prompts']

    mbox.showinfo( title='Welcome to the EHDN PubMed Parser', message= "\n".join(welcometext))
    while True:
        try:
            hd_list = ProcessPubmed()
            break
        except:
            if mbox.askokcancel('File Error', 'Do you want to select the PubMed file again'):
                continue
            else:
                quit()

    mbox.showinfo(title='Pubmed loaded', message='Please select the membership Excel file in the next file box')

    while True:
        try:
            data = ProcessMembers()
            break
        except:
            if mbox.askokcancel('File Error', 'Do you want to select the Excel file again'):
                continue
            else:
                quit()

    mbox.showinfo(title='Pubmed loaded', message='Please select the Pubmed CSV file in the next file box')

    while True:
        try:
            cites = ProcessCSV()
            break
        except:
            if mbox.askokcancel('File Error', 'Do you want to select the CSV file again'):
                continue
            else:
                quit()

    rpublist, apublist = CompareNamesN(data, cites)

    quarter = sd.askstring("Input", "What Quarter is this for? (i.e 1Q2025)")



    counter = 0
    document = Document()
    document.add_heading('Studies Conducted in Collaboration or in Affiliation with Regular EHDN Members in ' + quarter , level=1)
    for item in rpublist.values():
        counter += 1
        p = document.add_paragraph(str(counter) + '.\t')

        # item = item.split('\t')
        # namematch = item[0] + ' ' + item[1][0]
        for i in item.authors.split(', '):
            namematch = i.split(' ')[0] + ' ' + i.split(' ')[1][0]
            if any(namematch.lower() in sub for sub in item.name):
                p.add_run(i).bold = True
            elif any(normalize(namematch).lower() in sub for sub in item.name):
                p.add_run(i).bold = True
            else:
                p.add_run(i)
            if i != item.authors.split(',')[-1]:
                p.add_run(', ')
        p.add_run(' ' + item.title)
        p.add_run(' ' + item.citation)
        p.add_run(' ' + 'PMID:' + item.pmid)

    counter = 0
    document.add_heading('Studies Conducted in Collaboration or in Affiliation with Associate EHDN Members in ' + quarter, level=1)
    for item in apublist.values():
        counter += 1
        p = document.add_paragraph(str(counter) + '.\t')

        # item = item.split('\t')
        # namematch = item[0] + ' ' + item[1][0]
        for i in item.authors.split(', '):
            namematch = i.split(' ')[0] + ' ' + i.split(' ')[1][0]
            if any(namematch.lower() in sub for sub in item.name):
                p.add_run(i).bold = True
            elif any(normalize(namematch).lower() in sub for sub in item.name):
                p.add_run(i).bold = True
            else:
                p.add_run(i)
            if i != item.authors.split(',')[-1]:
                p.add_run(', ')
        p.add_run(' ' + item.title)
        p.add_run(' ' + item.citation)
        p.add_run(' ' + 'PMID:' + item.pmid)


    while True:
        try:
            pathout = fd.askdirectory(title='Please select the folder to save to')
            document.save(pathout + '/huntington_publications_' + quarter + '_' + datetime.today().strftime('%Y-%m-%d') + '.docx')
            break
        except:
            if mbox.askokcancel('Folder or File not accessible', 'Do you want to select the save folder again'):
                continue
            else:
                quit()


    mbox.showinfo( title='Complete', message= 'The output will be saved in the selected folder')





# See PyCharm help at https://www.jetbrains.com/help/pycharm/
